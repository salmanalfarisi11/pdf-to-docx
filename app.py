import re
import fitz                         # pip install PyMuPDF
from pdf2docx import Converter     # pip install pdf2docx
from docx import Document          # pip install python-docx
import tempfile, os, shutil, zipfile
import gradio as gr

# Regex pattern untuk URL
URL_PATTERN = re.compile(r'(https?://[^\s<>"\'\)\]]+)')

def convert_pdf_to_word(pdf_file) -> str:
    """
    Mengonversi satu PDF ke DOCX, menambahkan daftar link di akhir jika ada,
    lalu mengembalikan path ke file .docx.
    """
    workdir = tempfile.mkdtemp()
    try:
        # 1) Baca PDF & nama file asli
        if hasattr(pdf_file, 'read'):
            data = pdf_file.read()
            orig_name = getattr(pdf_file, 'name', 'output.pdf')
        elif isinstance(pdf_file, dict) and 'name' in pdf_file:
            with open(pdf_file['name'], 'rb') as f:
                data = f.read()
            orig_name = pdf_file['name']
        elif isinstance(pdf_file, str):
            with open(pdf_file, 'rb') as f:
                data = f.read()
            orig_name = pdf_file
        else:
            raise ValueError("Unsupported input type")
        base_name = os.path.splitext(os.path.basename(orig_name))[0]

        # 2) Simpan PDF sementara
        pdf_path = os.path.join(workdir, 'input.pdf')
        with open(pdf_path, 'wb') as f:
            f.write(data)

        # 3) Konversi ke DOCX
        temp_docx = os.path.join(workdir, 'output.docx')
        cv = Converter(pdf_path)
        cv.convert(temp_docx, start=0, end=None)
        cv.close()

        # 4) Format semua tabel
        doc = Document(temp_docx)
        for table in doc.tables:
            table.style = 'Table Grid'
        doc.save(temp_docx)

        # 5) Ekstrak semua link (annotation + regex)
        links = []
        pdf_doc = fitz.open(pdf_path)
        for page in pdf_doc:
            for annot in page.annots() or []:
                uri = annot.info.get('uri')
                if uri:
                    u = uri.rstrip('.,;:)]')
                    if u not in links:
                        links.append(u)
            text = page.get_text('text')
            for m in URL_PATTERN.findall(text):
                u = m.rstrip('.,;:)]')
                if u not in links:
                    links.append(u)
        pdf_doc.close()

        # 6) Tambahkan daftar link di DOCX jika ada
        if links:
            doc = Document(temp_docx)
            doc.add_page_break()
            doc.add_heading('Daftar Link', level=2)
            for u in links:
                doc.add_paragraph(u)
            doc.save(temp_docx)

        # 7) Salin ke file akhir dengan nama asli
        final_path = os.path.join(workdir, f"{base_name}.docx")
        shutil.copy(temp_docx, final_path)
        return final_path

    finally:
        # jangan dihapus agar Gradio masih bisa mengakses file
        pass

def convert_and_enable(pdf_files):
    """
    Menerima list PDF, memproses semuanya,
    lalu mengembalikan gr.update() dengan path .docx tunggal
    atau .zip jika lebih dari satu file.
    """
    out_paths = [convert_pdf_to_word(pdf) for pdf in pdf_files]
    # Jika >1, bungkus ke ZIP
    if len(out_paths) > 1:
        zip_dir = tempfile.mkdtemp()
        zip_path = os.path.join(zip_dir, "converted_docs.zip")
        with zipfile.ZipFile(zip_path, "w") as zf:
            for p in out_paths:
                zf.write(p, arcname=os.path.basename(p))
        return gr.update(value=zip_path, interactive=True)
    # Jika hanya satu, langsung .docx
    return gr.update(value=out_paths[0], interactive=True)

def reset_download(_):
    """
    Reset DownloadButton ke state awal ketika input berubah/clear.
    """
    return gr.update(value=None, interactive=False)

# CSS untuk tombol full-width
css = """
#convert-btn, #download-btn {
    width: 100%;
}
"""

with gr.Blocks(css=css, title="PDF→Word Converter") as demo:
    gr.Markdown("# PDF→Word Converter 🎉")
    gr.Markdown("Upload satu atau lebih PDF, lalu tekan Convert untuk mendapatkan DOCX atau ZIP.")

    # Upload multiple PDFs
    pdf_inputs = gr.Files(
        label="Upload PDF(s)",
        file_types=['.pdf']
    )

    # Tombol Convert
    convert_btn = gr.Button(
        "Convert",
        variant="primary",
        elem_id="convert-btn"
    )

    # Tombol Download (disabled hingga ready)
    download_btn = gr.DownloadButton(
        label="⬇️ Download Output",
        value=None,
        interactive=False,
        variant="primary",
        elem_id="download-btn"
    )

    # Reset download button setiap kali input berubah atau di-clear
    pdf_inputs.change(
        fn=reset_download,
        inputs=pdf_inputs,
        outputs=download_btn
    )

    # Proses convert saat tombol ditekan, lalu enable download
    convert_btn.click(
        fn=convert_and_enable,
        inputs=pdf_inputs,
        outputs=download_btn
    )

    gr.Markdown("---\nBuilt with ❤️ using Gradio dan PyMuPDF.")

if __name__ == "__main__":
    demo.launch(share=True)
